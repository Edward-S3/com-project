#!/usr/bin/env python3
"""仕様書・ユーザーマニュアル（docx）を生成する。"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

OUTPUT_DIR = Path("/opt/exam/docs")
SPEC_PATH = OUTPUT_DIR / "試験システム_仕様書.docx"
MANUAL_PATH = OUTPUT_DIR / "試験システム_ユーザーマニュアル.docx"
TODAY = date.today().strftime("%Y年%m月%d日")


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    style.font.size = Pt(10.5)


def add_title(doc: Document, text: str, subtitle: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11)
    doc.add_paragraph()


def add_h(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str):
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build_specification():
    doc = Document()
    set_doc_defaults(doc)
    add_title(
        doc,
        "試験問題作成・採点・評価システム",
        "ソフトウェア仕様書",
    )
    add_p(doc, f"文書版: 1.0　作成日: {TODAY}")
    add_p(doc, "対象: システム管理者・開発者・運用担当者")

    add_h(doc, "1. 概要", 1)
    add_p(
        doc,
        "本システムは、社内・組織内向けの Web 試験プラットフォームである。"
        "問題作成者が試験（設問）を登録し、受験者は URL からログイン不要で受験する。"
        "自動採点（選択式・記述式 AI 採点）、結果メール送信、受験結果の分析レポート出力を提供する。",
    )
    add_bullets(
        doc,
        [
            "アプリケーション名（画面タイトル）: 試験問題作成・採点・評価システム",
            "実装: Python 3 + Streamlit（単一ファイル exam_app.py）",
            "データベース: SQLite（exam_app.db）",
            "既定アクセス: http://{ホスト}/exam/（受験は ?ID=試験コード）",
        ],
    )

    add_h(doc, "2. システム構成", 1)
    add_h(doc, "2.1 技術スタック", 2)
    add_table(
        doc,
        ["区分", "技術・製品"],
        [
            ["UI", "Streamlit（wide レイアウト）"],
            ["DB", "SQLite3"],
            ["データ処理", "pandas, json"],
            ["グラフ", "matplotlib（日本語フォント IPAex Gothic）"],
            ["AI 採点・分析", "Google Gemini API（gemini-2.5-flash、任意）"],
            ["PDF レポート", "FPDF（任意）"],
            ["メール", "SMTP（smtplib、HTML メール）"],
            ["設定", "python-dotenv（.env）"],
            ["プロセス管理", "systemd（exam-app.service）"],
        ],
    )

    add_h(doc, "2.2 画面ルーティング", 2)
    add_bullets(
        doc,
        [
            "クエリパラメータ ID あり → 受験者画面（認証不要）",
            "上記以外 → ログイン画面 → ロールに応じた管理／作成者ダッシュボード",
        ],
    )

    add_h(doc, "3. データモデル", 1)
    add_h(doc, "3.1 テーブル一覧", 2)
    add_table(
        doc,
        ["テーブル", "説明"],
        [
            ["users", "管理者・問題作成者アカウント"],
            ["exams", "試験定義（設問 schema、評価設定 grading_config）"],
            ["submissions", "受験者の回答・採点結果"],
            ["exam_editors", "試験ごとの編集・分析権限の委譲"],
        ],
    )

    add_h(doc, "3.2 users", 2)
    add_table(
        doc,
        ["カラム", "型", "説明"],
        [
            ["id", "INTEGER PK", "ユーザー ID"],
            ["username", "TEXT UNIQUE", "ログイン ID"],
            ["password", "TEXT", "SHA-256 ハッシュ"],
            ["company_name", "TEXT", "所属名・組織名"],
            ["role", "TEXT", "admin または creator"],
            ["created_at", "DATETIME", "登録日時"],
        ],
    )

    add_h(doc, "3.3 exams", 2)
    add_table(
        doc,
        ["カラム", "型", "説明"],
        [
            ["id", "TEXT PK", "試験コード（例 EX-A1B2C3）"],
            ["title", "TEXT", "試験タイトル"],
            ["description", "TEXT", "説明・概要"],
            ["status", "TEXT", "公開状態（新規作成時は「公開」固定）"],
            ["limit_time", "INTEGER", "制限時間（分、0=無制限）"],
            ["schema", "TEXT", "設問 JSON 配列"],
            ["grading_config", "TEXT", "総合評価設定 JSON"],
            ["created_by", "INTEGER FK", "作成者 users.id"],
            ["created_at", "DATETIME", "作成日時"],
        ],
    )

    add_h(doc, "3.4 設問オブジェクト（schema 配列要素）", 2)
    add_table(
        doc,
        ["フィールド", "説明"],
        [
            ["id", "設問 ID（例 q1）"],
            ["type", "択一選択 / 複数選択 / ○×式 / テキスト（記述式） / テキストエリア（長文記述）"],
            ["question", "問題文"],
            ["category", "カテゴリ（分野）"],
            ["points", "配点（数値）"],
            ["options", "選択肢配列（選択式のみ）"],
            ["correct_answer", "正解（文字列または配列）"],
            ["explanation", "解説文（採点結果表示用）"],
            ["option_explanations", "選択肢ごとの解説（辞書、任意）"],
        ],
    )

    add_h(doc, "3.5 grading_config（総合評価）", 2)
    add_p(doc, "試験全体の得点に対する表示ラベルを定義する JSON。")
    add_table(
        doc,
        ["mode", "内容"],
        [
            ["none", "総合評価を表示しない（得点のみ）"],
            ["pass_fail", "A パターン: 1 基準で合格／不合格ラベル"],
            ["tiers", "B パターン: 複数基準点と段階ラベル（上から高い基準で判定）"],
        ],
    )
    add_p(doc, "score_type: absolute（得点）または percent（正解率％）。ラベル文字列は任意（例 ARank、Pass）。")

    add_h(doc, "3.6 submissions", 2)
    add_table(
        doc,
        ["カラム", "説明"],
        [
            ["exam_id", "試験 ID"],
            ["examinee_name", "受験者氏名"],
            ["examinee_email", "メールアドレス"],
            ["answers", "回答 JSON（設問 ID → 回答）"],
            ["score", "合計得点"],
            ["total_points", "満点"],
            ["results", "設問別採点結果 JSON"],
            ["email_sent", "結果メール送信成否（0/1）"],
            ["submitted_at", "提出日時"],
        ],
    )

    add_h(doc, "4. 権限・ロール", 1)
    add_table(
        doc,
        ["ロール", "識別", "主な権限"],
        [
            ["管理者", "admin", "全ユーザー管理、全試験・全受験結果の閲覧・分析"],
            ["問題作成者", "creator", "自試験の作成・編集、権限付与試験の編集・分析"],
            ["受験者", "（アカウントなし）", "URL からの受験・結果閲覧のみ"],
        ],
    )
    add_p(
        doc,
        "exam_editors: 試験オーナーが他の creator/admin に編集・結果分析権限を付与。"
        "付与先は設問編集および「担当試験の受験結果・分析」で当該試験を参照可能。",
    )

    add_h(doc, "5. 機能仕様", 1)

    add_h(doc, "5.1 認証", 2)
    add_bullets(
        doc,
        [
            "ログイン画面: ユーザー名・パスワード",
            "パスワードは SHA-256 で保存（ソルトなし）",
            "セッション: Streamlit session_state（logged_in, role, user_id 等）",
            "初期アカウント（DB 空時）: admin/admin123, creator1/creator123",
        ],
    )

    add_h(doc, "5.2 試験管理（作成者）", 2)
    add_bullets(
        doc,
        [
            "新規試験: タイトル、説明、試験コード（自動生成）、制限時間、設問、評価設定",
            "設問ビルダー: 5 形式、配点・解説・選択肢解説",
            "CSV 取り込み: UTF-8 / Shift_JIS、置換または末尾追加",
            "既存試験編集・コピー新規・編集・分析権限の付与",
            "受験 URL 表示・クリップボードコピー（EXAM_HOST/PORT から生成）",
        ],
    )

    add_h(doc, "5.3 採点ロジック", 2)
    add_table(
        doc,
        ["設問形式", "採点方法"],
        [
            ["択一選択・○×式", "正解文字列との完全一致"],
            ["複数選択", "正解集合と回答集合の一致"],
            ["記述式・長文", "Gemini API による意味評価（未設定時は簡易文字列判定）"],
        ],
    )
    add_p(doc, "記述式の is_correct は得点が配点の 60% 以上で true。")

    add_h(doc, "5.4 受験者フロー", 2)
    add_bullets(
        doc,
        [
            "状態: register → in_progress → completed",
            "登録: 氏名、メール（必須・簡易検証）",
            "制限時間: 残り時間表示（0=無制限）",
            "提出: 全設問採点後 DB 保存、SMTP 設定時は HTML メール送信",
            "完了画面: 得点、正解率、総合評価、設問別解説",
        ],
    )

    add_h(doc, "5.5 分析・レポート", 2)
    add_bullets(
        doc,
        [
            "基本統計: 受験者数、平均・最高・最低得点、評価内訳",
            "受験者明細一覧（行選択で詳細・メール再送信）",
            "設問別正答率グラフ",
            "AI 全体傾向分析（Gemini、セッションキャッシュ）",
            "PDF 分析レポート・CSV ダウンロード（Shift_JIS / UTF-8）",
        ],
    )

    add_h(doc, "5.6 メール", 2)
    add_p(doc, ".env の SMTP_* / EMAIL_FROM が有効な場合のみ送信。プレースホルダ値は未設定扱い。")

    add_h(doc, "6. 外部連携・環境変数", 1)
    add_table(
        doc,
        ["変数名", "用途"],
        [
            ["GEMINI_API_KEY", "AI 採点・全体分析"],
            ["SMTP_SERVER / SMTP_PORT", "メールサーバ"],
            ["SMTP_USER / SMTP_PASSWORD", "SMTP 認証"],
            ["EMAIL_FROM", "送信元アドレス"],
            ["EXAM_HOST", "受験 URL のホスト（既定 172.16.16.10）"],
            ["EXAM_PORT", "受験 URL のポート（既定 8505）"],
        ],
    )

    add_h(doc, "7. デプロイメント", 1)
    add_bullets(
        doc,
        [
            "作業ディレクトリ: /opt/exam",
            "仮想環境: /opt/exam/venv",
            "systemd ユニット: exam-app.service",
            "起動: streamlit run exam_app.py --server.port=8505 --server.address=0.0.0.0",
            "再起動: sudo systemctl restart exam-app",
        ],
    )

    add_h(doc, "8. 制約・既知の仕様", 1)
    add_bullets(
        doc,
        [
            "試験 status の下書き機能は UI 上なし（常に公開で登録）",
            "同一メールでの再受験は可能（最新提出を結果画面に表示）",
            "制限時間超過時は警告表示するが、提出はフォームの「回答を提出する」操作が必要",
            "fpdf / google-generativeai 未インストール時は該当機能を無効化して継続動作",
        ],
    )

    add_h(doc, "9. ファイル構成", 1)
    add_table(
        doc,
        ["パス", "説明"],
        [
            ["/opt/exam/exam_app.py", "アプリケーション本体"],
            ["/opt/exam/exam_app.db", "SQLite データ"],
            ["/opt/exam/.env", "環境変数"],
            ["/opt/exam/.streamlit/config.toml", "Streamlit 設定"],
            ["/opt/exam/exam-app.service", "systemd ユニット定義"],
            ["/opt/exam/requirements.txt", "Python 依存パッケージ"],
        ],
    )

    doc.save(SPEC_PATH)
    print(f"Created: {SPEC_PATH}")


def build_user_manual():
    doc = Document()
    set_doc_defaults(doc)
    add_title(
        doc,
        "試験問題作成・採点・評価システム",
        "ユーザーマニュアル",
    )
    add_p(doc, f"文書版: 1.0　作成日: {TODAY}")
    add_p(doc, "対象: 管理者・問題作成者・受験者")

    add_h(doc, "はじめに", 1)
    add_p(
        doc,
        "本マニュアルは、試験の作成・配布・受験・結果確認の手順を説明します。"
        "ブラウザでシステムにアクセスし、役割に応じた操作を行ってください。",
    )
    add_table(
        doc,
        ["役割", "主な作業"],
        [
            ["管理者", "ユーザー登録、全試験・全結果の管理"],
            ["問題作成者", "試験作成、受験 URL 配布、担当試験の結果分析"],
            ["受験者", "URL から受験・結果確認（ログイン不要）"],
        ],
    )

    add_h(doc, "第1章　共通操作", 1)

    add_h(doc, "1.1 アクセス方法", 2)
    add_bullets(
        doc,
        [
            "管理・作成画面: http://（サーバーアドレス）:8505/",
            "受験画面: http://（サーバーアドレス）/exam/?ID=（試験コード）",
            "推奨ブラウザ: Chrome、Edge、Firefox など最新版",
        ],
    )

    add_h(doc, "1.2 ログイン（管理者・問題作成者）", 2)
    add_p(doc, "画面タイトル「🔐 管理者・作成者ログイン」で次を入力します。")
    add_table(
        doc,
        ["項目", "入力内容"],
        [
            ["ユーザー名", "付与されたログイン ID"],
            ["パスワード", "初期パスワードまたは変更後のパスワード"],
        ],
    )
    add_p(doc, "「ログイン」ボタンを押すと、権限に応じたメニューが表示されます。")
    add_p(doc, "ログアウトはサイドバー下部の「ログアウト」です。")

    add_h(doc, "1.3 パスワード・所属の変更", 2)
    add_p(doc, "サイドバー「自身のパスワード・所属の変更」から次を変更できます。")
    add_bullets(
        doc,
        [
            "所属名・組織名",
            "パスワード（現在のパスワード入力が必須）",
        ],
    )
    add_p(doc, "「設定を更新する」で保存します。")

    add_h(doc, "第2章　問題作成者向け操作", 1)

    add_h(doc, "2.1 メニュー一覧", 2)
    add_table(
        doc,
        ["メニュー", "内容"],
        [
            ["試験問題の作成・編集", "試験の新規作成、設問編集、URL 確認、権限付与"],
            ["担当試験の受験結果・分析", "自分が作成した試験、または権限を付与された試験の分析"],
            ["自身のパスワード・所属の変更", "プロフィール変更"],
        ],
    )

    add_h(doc, "2.2 新規試験の作成", 2)
    add_p(doc, "「試験問題の作成・編集」→「新規試験の作成」で以下を設定します。")
    add_table(
        doc,
        ["画面項目", "説明"],
        [
            ["試験タイトル", "必須。受験者にも表示されます"],
            ["試験の説明・概要", "任意。受験開始画面に表示"],
            ["試験コード", "自動生成（EX-XXXXXX）。「試験コードを再生成」で変更可"],
            ["制限時間（分）", "0 で無制限。1〜180 分"],
            ["総合評価・合格判定の設定", "後述"],
            ["設問の設定", "手入力または CSV 取り込み"],
        ],
    )
    add_p(doc, "設問を 1 問以上登録し、「試験を登録・公開する」を押すと公開されます。")
    add_p(doc, "登録後に表示される「受験者への配布用URL」を受験者に共有してください。")

    add_h(doc, "2.3 設問の追加方法", 2)
    add_h(doc, "（1）画面から追加", 3)
    add_bullets(
        doc,
        [
            "「追加する設問形式」で形式を選び「この形式で設問を追加」",
            "問題文・配点・正解・解説などを入力",
            "不要な設問は「削除」で削除",
        ],
    )
    add_h(doc, "（2）CSV から取り込み", 3)
    add_bullets(
        doc,
        [
            "「CSVから設問を取り込む」を開く",
            "テンプレートをダウンロード（Shift_JIS または UTF-8）",
            "CSV を選択し「取り込み方法」（置き換え／末尾追加）を選ぶ",
            "「CSVを取り込む」を実行",
        ],
    )
    add_p(doc, "CSV の主な列: 設問形式、問題文、カテゴリ、配点、選択肢、正解、問題解説、選択肢解説")

    add_h(doc, "2.4 総合評価・合格判定の設定", 2)
    add_table(
        doc,
        ["評価方式", "用途"],
        [
            ["なし", "得点のみ表示"],
            ["Aパターン：合格／不合格", "1 つの基準点で 2 種類のラベル（例: 合格 / 不合格）"],
            ["Bパターン：段階評価", "複数基準点とラベル（段数・数値・表示名はすべて任意）"],
        ],
    )
    add_p(doc, "基準の単位は「得点（点）」または「正解率（％）」を選択できます。")
    add_p(doc, "B パターンでは「基準点を追加」「削除」で段数を変更し、各段の数値と表示ラベル（例: ARank）を入力します。")
    add_p(doc, "「サンプル（90/80/75/60/0点）を読み込む」は初期例です。読み込み後に自由に編集してください。")

    add_h(doc, "2.5 既存試験の編集・コピー", 2)
    add_bullets(
        doc,
        [
            "一覧の展開から「編集」→ 内容変更 →「変更を保存する」",
            "「コピーして新規作成」で別コードの新試験として複製",
            "編集をやめる場合は「編集をキャンセル」",
        ],
    )

    add_h(doc, "2.6 編集・分析権限の付与", 2)
    add_p(doc, "試験の作成者のみ、他の問題作成者に権限を付与できます。")
    add_bullets(
        doc,
        [
            "「編集・結果分析権限の付与」でユーザーを選択",
            "「編集・分析権限を付与」で追加",
            "「解除」で権限を削除",
            "付与先は設問編集と「担当試験の受験結果・分析」が可能",
        ],
    )

    add_h(doc, "2.7 受験結果・分析の見方", 2)
    add_p(doc, "「担当試験の受験結果・分析」で試験を選択します。")
    add_bullets(
        doc,
        [
            "基本統計: 受験者数、平均・最高・最低得点",
            "受験者ごとの回答明細: 表の行をクリックで詳細表示",
            "設問別正答率のグラフ",
            "AI による解答傾向の総合分析（API 設定時）",
            "PDF レポート・CSV のダウンロード",
        ],
    )
    add_p(doc, "CSV の文字コードは Shift_JIS または UTF-8 を選択できます。")
    add_p(doc, "メール再送信: 明細で受験者を選択し「採点結果メールを再送信」（SMTP 設定時）")

    add_h(doc, "第3章　管理者向け操作", 1)

    add_h(doc, "3.1 メニュー一覧", 2)
    add_table(
        doc,
        ["メニュー", "内容"],
        [
            ["ユーザー(他者)の登録と管理", "問題作成者の登録・権限変更・削除"],
            ["すべての試験問題管理", "全試験の閲覧・編集・コピー"],
            ["すべての受験結果・分析", "任意の試験の分析"],
            ["自身のパスワード・所属の変更", "プロフィール変更"],
        ],
    )

    add_h(doc, "3.2 ユーザーの登録", 2)
    add_p(doc, "「新規問題作成ユーザー（他者）の登録」でログイン ID、所属名、パスワードを入力し「ユーザーを登録する」。")
    add_p(doc, "新規ユーザーは「問題作成者（creator）」として登録されます。")

    add_h(doc, "3.3 ユーザーの編集・削除", 2)
    add_bullets(
        doc,
        [
            "一覧からユーザーを選択",
            "所属名・権限（一般問題作成者 / システム管理者）・パスワードを変更",
            "削除は展開「このユーザーを削除する」で承諾チェック後に実行",
        ],
    )

    add_h(doc, "3.4 試験・結果の管理", 2)
    add_p(doc, "全試験を一覧表示し、作成者名付きで編集・コピー・受験 URL 確認ができます。")
    add_p(doc, "「すべての受験結果・分析」は問題作成者と同様の分析機能を全試験に適用します。")

    add_h(doc, "第4章　受験者向け操作", 1)

    add_h(doc, "4.1 受験の流れ", 2)
    add_table(
        doc,
        ["手順", "操作"],
        [
            ["1", "配布された URL をブラウザで開く"],
            ["2", "試験説明を確認し、氏名・メールアドレスを入力"],
            ["3", "「試験を開始する」"],
            ["4", "各設問に回答（制限時間がある場合は残り時間を確認）"],
            ["5", "「回答を提出する」"],
            ["6", "採点結果・設問別解説を確認"],
        ],
    )

    add_h(doc, "4.2 受験時の注意", 2)
    add_bullets(
        doc,
        [
            "メールアドレスは結果通知用（SMTP 設定時）および本人確認に使用",
            "制限時間がある試験は時間内に「回答を提出する」こと",
            "提出後、得点・正解率・総合評価（設定がある場合）が表示される",
            "選択式で選んだ選択肢に解説がある場合、結果画面に表示される",
            "「トップ画面に戻る」で URL パラメータがクリアされる",
        ],
    )

    add_h(doc, "4.3 結果メール", 2)
    add_p(
        doc,
        "サーバーでメール送信が設定されている場合、提出後に登録メールへ採点結果が送られます。"
        "届かない場合は迷惑メールフォルダを確認し、画面の結果も必ず確認してください。",
    )

    add_h(doc, "第5章　よくある質問（FAQ）", 1)
    add_table(
        doc,
        ["質問", "回答"],
        [
            ["受験 URL が開けない", "試験コード（ID）が正しいか、サーバーが起動しているか管理者に確認"],
            ["ログインできない", "ユーザー名・パスワード、Caps Lock を確認。管理者に再発行依頼"],
            ["編集権限がない", "試験作成者または権限付与を依頼"],
            ["メールが届かない", "SMTP 未設定の可能性。画面で結果確認。管理者に問い合わせ"],
            ["PDF がダウンロードできない", "サーバーに PDF ライブラリ未導入の可能性。CSV は利用可"],
            ["記述式の点数が期待と違う", "AI 採点のため模範解答・キーワードを明確に記載"],
        ],
    )

    add_h(doc, "第6章　お問い合わせ", 1)
    add_p(doc, "システムの障害・権限・メール設定・受験 URL は、貴組織のシステム管理者へご連絡ください。")

    doc.save(MANUAL_PATH)
    print(f"Created: {MANUAL_PATH}")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_specification()
    build_user_manual()


if __name__ == "__main__":
    main()
